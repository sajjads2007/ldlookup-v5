# v5.0.1.2 — LD Lookup
# - Oxford blue theme + white title + disclaimer
# - Clean table ("View" link instead of long URL)
# - Thumbnails + in-app full-size preview via st.modal
# - Excel & PDF exports (optional image embedding)

import io
import re
from pathlib import Path
from typing import List, Optional

import pandas as pd
import requests
import streamlit as st

# Optional readers for DOCX/PDF
import pdfplumber              # PDF text
from docx import Document      # DOCX text

# -------------------------------------------------
# Page & Styling
# -------------------------------------------------
OXFORD = "#0B132B"
st.set_page_config(page_title="LD Lookup", page_icon="🧾", layout="wide")

st.markdown(
    f"""
    <style>
      .stApp, .main {{
        background-color: {OXFORD};
      }}
      .block-container {{
        background: #ffffff;
        border-radius: 16px;
        padding: 18px 18px 8px 18px;
        box-shadow: 0 2px 12px rgba(0,0,0,.12);
      }}
      .topbar {{
        background: {OXFORD};
        color: #fff;
        padding: 14px 18px;
        border-radius: 12px;
        margin-bottom: 10px;
      }}
      .topbar h1 {{ margin: 0; font-size: 22px; color: #fff; }}
      .disclaimer {{ color:#C9CED8; font-size:.92rem; margin:4px 0 12px 0; }}
      .runrow {{ position: sticky; top: 0; z-index: 5; background: white; padding: 6px 0 2px; }}
      .thumb {{
        height: 70px; width: auto; object-fit: contain; border-radius: 6px;
        border: 1px solid #e6e6e6; background: #fff;
      }}
      .tbl table {{ border-collapse: collapse; width: 100%; }}
      .tbl th, .tbl td {{ border: 1px solid #e6e6e6; padding: 8px; vertical-align: middle; font-size: .92rem; }}
      .tbl th {{ background: #f6f7fb; color: {OXFORD}; text-align: left; }}
      @media (max-width: 640px) {{
        .thumb {{ height: 60px; }}
        .tbl th, .tbl td {{ font-size: .88rem; }}
      }}
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    """
    <div class="topbar">
      <h1>LD Lookup — Version 5.0</h1>
      <div class="disclaimer">This is not an official app — currently in debugging mode</div>
    </div>
    """,
    unsafe_allow_html=True,
)

# -------------------------------------------------
# Config & Helpers
# -------------------------------------------------
IMG_WIDTH = 640
IMG_QUALITY = 75
CDN_TEMPLATE = "https://cdn-tp2.mozu.com/28945-m4/cms/files/{L}.jpg?w={w}&q={q}"
LNUM_RE = re.compile(r"\bL\d+\b", flags=re.IGNORECASE)

def extract_lnumbers_from_text(text: str) -> List[str]:
    if not text:
        return []
    candidates = [m.group(0).upper() for m in LNUM_RE.finditer(text)]
    seen, out = set(), []
    for c in candidates:
        if c not in seen:
            seen.add(c); out.append(c)
    return out

def extract_lnumbers_from_dataframe(df: pd.DataFrame) -> List[str]:
    lnumbers, seen = [], set()
    for col in df.columns:
        for val in df[col].astype(str).tolist():
            for l in extract_lnumbers_from_text(val):
                if l not in seen:
                    seen.add(l); lnumbers.append(l)
    return lnumbers

def build_image_url(l_number: str, width=IMG_WIDTH, quality=IMG_QUALITY) -> str:
    l = str(l_number).strip().upper()
    return CDN_TEMPLATE.format(L=l, w=width, q=quality)

def read_any_file(file) -> str:
    """
    Return concatenated text from uploaded file.
    Supports: CSV, XLSX/XLS, TXT, DOCX, PDF.
    """
    name = file.name
    suffix = Path(name).suffix.lower()

    if suffix in [".xlsx", ".xls", ".csv"]:
        try:
            if suffix == ".csv":
                df = pd.read_csv(file, dtype=str, on_bad_lines="skip")
            else:
                df = pd.read_excel(file, dtype=str)
            return "\n".join(df.astype(str).fillna("").agg(" ".join, axis=1).tolist())
        except Exception as e:
            return f"READ_ERROR: {e}"

    if suffix == ".txt":
        return file.read().decode("utf-8", errors="ignore")

    if suffix == ".docx":
        try:
            doc = Document(file)
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except Exception as e:
            return f"READ_ERROR: {e}"

    if suffix == ".pdf":
        try:
            with pdfplumber.open(file) as pdf:
                texts = []
                for page in pdf.pages:
                    texts.append(page.extract_text() or "")
                return "\n".join(texts)
        except Exception as e:
            return f"READ_ERROR: {e}"

    if suffix == ".doc":
        return "READ_ERROR: Legacy .doc detected. Convert to .docx or PDF and re-upload."

    return "READ_ERROR: Unsupported file type."

def make_excel_bytes(df: pd.DataFrame) -> bytes:
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
        df_xlsx = df.copy()

        def excel_img(url, alt_text=""):
            safe_url = str(url).replace('"', '""')
            safe_alt = str(alt_text).replace('"', '""')
            return f'=IMAGE("{safe_url}","{safe_alt}")'

        df_xlsx["Photo"] = [excel_img(u, l) for u, l in zip(df_xlsx["ImageURL"], df_xlsx["LNumber"])]
        df_xlsx[["LNumber", "ImageURL", "Photo"]].to_excel(writer, index=False, sheet_name="Results")
        ws = writer.sheets["Results"]
        ws.set_column(0, 0, 14)  # LNumber
        ws.set_column(1, 1, 60)  # URL
        ws.set_column(2, 2, 18)  # Photo
    out.seek(0)
    return out.read()

def make_pdf_bytes(df: pd.DataFrame, include_images: bool = False, thumb_h: int = 72) -> bytes:
    """
    Lightweight PDF export. If include_images=True, downloads images and embeds thumbnails.
    """
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.utils import ImageReader

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), leftMargin=24, rightMargin=24, topMargin=24, bottomMargin=24)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("LD Lookup — Results", styles["Title"]))
    story.append(Paragraph("Version 5.0.1.2", styles["Normal"]))
    story.append(Spacer(1, 10))

    data = [["LNumber", "ImageURL", "Photo (thumb)"]]
    for _, r in df.iterrows():
        lnum = str(r["LNumber"])
        url  = str(r["ImageURL"])
        cell_img = ""
        if include_images:
            try:
                resp = requests.get(url, timeout=5)
                if resp.status_code // 100 == 2:
                    ir = ImageReader(io.BytesIO(resp.content))
                    img = Image(ir, hAlign="LEFT")
                    iw, ih = ir.getSize()
                    scale = thumb_h / float(ih)
                    img.drawHeight = thumb_h
                    img.drawWidth = iw * scale
                    cell_img = img
            except Exception:
                cell_img = ""
        data.append([lnum, url, cell_img])

    tbl = Table(data, colWidths=[90, 360, 120])
    tbl.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("TEXTCOLOR", (0,0), (-1,0), colors.HexColor(OXFORD)),
        ("BACKGROUND", (0,0), (-1,0), colors.whitesmoke),
        ("GRID", (0,0), (-1,-1), 0.5, colors.lightgrey),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F9FAFB")]),
    ]))
    story.append(tbl)
    doc.build(story)
    buf.seek(0)
    return buf.read()

def maybe_lookup_name_for_lnumber(l_number: str) -> Optional[str]:
    # Placeholder: wire to a real source if/when available.
    return None

# -------------------------------------------------
# Inputs
# -------------------------------------------------
with st.expander("Input options", expanded=True):
    c1, c2 = st.columns([1.3, 1])
    with c1:
        uploaded = st.file_uploader(
            "Upload CSV / XLSX / XLS / TXT / PDF / DOCX",
            type=["csv", "xlsx", "xls", "txt", "pdf", "docx"],
            key="file_input_v5_modal"
        )
        st.caption("Tip: For old .doc, convert to .docx first.")
    with c2:
        pasted = st.text_area("Or paste any text", height=180, key="pasted_v5_modal")

st.markdown('<div class="runrow"></div>', unsafe_allow_html=True)
run = st.button("🚀 Run L-Number Lookup", use_container_width=True, type="primary")
if not run:
    st.stop()

# -------------------------------------------------
# Collect data
# -------------------------------------------------
lnumbers: List[str] = []

if uploaded is not None:
    text = read_any_file(uploaded)
    if text.startswith("READ_ERROR:"):
        st.error(text); st.stop()
    lnumbers = extract_lnumbers_from_text(text)

if not lnumbers and pasted.strip():
    lnumbers = extract_lnumbers_from_text(pasted)

if not lnumbers:
    st.warning("No L-numbers detected. Provide data containing items like **L1304179**.")
    st.stop()

# Build results dataframe
df = pd.DataFrame({"LNumber": lnumbers})
df["LNumber"] = df["LNumber"].str.upper().str.strip()
df["ImageURL"] = df["LNumber"].apply(build_image_url)
# Optional names (left blank if not available)
df["Name"] = [maybe_lookup_name_for_lnumber(l) or "" for l in df["LNumber"]]

# -------------------------------------------------
# Table (clean links + thumbnails)
# -------------------------------------------------
def link_text(url: str, label: str = "View") -> str:
    return f'<a href="{url}" target="_blank">{label}</a>'

rows_html = []
for _, r in df.iterrows():
    lnum = r["LNumber"]
    url  = r["ImageURL"]
    name = r["Name"]
    thumb = f'<img class="thumb" src="{url}" alt="{lnum}"/>'
    link  = link_text(url, "View")
    rows_html.append(
        f"<tr><td>{lnum}</td><td>{name}</td><td>{link}</td><td>{thumb}</td></tr>"
    )

st.subheader("Results")
st.markdown(
    f"""
    <div class="tbl">
      <table>
        <thead><tr><th>LNumber</th><th>Name</th><th>Image</th><th>Photo (thumb)</th></tr></thead>
        <tbody>{''.join(rows_html)}</tbody>
      </table>
    </div>
    """,
    unsafe_allow_html=True,
)

# -------------------------------------------------
# In-app full-size preview (modal)
# -------------------------------------------------
st.markdown("**Tap a thumbnail below to preview in-app:**")

PREVIEW_COLS = st.slider("Preview grid columns", 3, 10, 6, key="grid_cols_v5")
grid = st.columns(PREVIEW_COLS)

# keep selected url in session
if "preview_url" in st.session_state and st.session_state["preview_url"]:
    with st.modal("Preview", key="img_modal"):
        st.image(st.session_state["preview_url"], use_column_width=True)
        st.caption(st.session_state.get("preview_caption", ""))

for i, r in df.iterrows():
    with grid[i % PREVIEW_COLS]:
        st.image(r["ImageURL"], caption=r["LNumber"], use_container_width=True)
        if st.button("🔍 Preview", key=f"pv_{i}"):
            st.session_state["preview_url"] = r["ImageURL"]
            st.session_state["preview_caption"] = r["LNumber"]
            st.rerun()

# -------------------------------------------------
# Exports
# -------------------------------------------------
st.divider()
st.subheader("Export")

c1, c2, _ = st.columns([1, 1, 2])
with c1:
    st.download_button(
        "⬇️ Excel (with IMAGE formula)",
        data=make_excel_bytes(df),
        file_name="ld_lookup_v5.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

with c2:
    include_pics = st.toggle("Embed images in PDF (slower)", value=False, help="Downloads thumbnails into the PDF.")
    st.download_button(
        "⬇️ PDF",
        data=make_pdf_bytes(df, include_images=include_pics),
        file_name="ld_lookup_v5.pdf",
        mime="application/pdf",
    )
